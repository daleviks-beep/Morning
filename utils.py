import io
import json
import time
import tempfile
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

import pdfplumber
import requests
from docx import Document


# ============================================================
# CONFIG
# ============================================================

OPENAI_RESPONSES_URL = "https://api.openai.com/v1/responses"
GAMMA_FROM_TEMPLATE_URL = "https://public-api.gamma.app/v1.0/generations/from-template"
GAMMA_STATUS_URL = "https://public-api.gamma.app/v1.0/generations/{generation_id}"


@dataclass
class AppConfig:
    openai_api_key: str
    gamma_api_key: str
    openai_model: str = "gpt-5"
    gamma_id: str = "g_pepbfk69p9lagj1"
    gamma_theme_id: str = "ge1kywkagyzapfv"
    gamma_folder_ids: list | None = None


# ============================================================
# SMALL HELPERS
# ============================================================

def mask_key(key: str) -> str:
    if not key:
        return "not set"
    if len(key) <= 8:
        return "*" * len(key)
    return f"{key[:4]}...{key[-4:]}"


def request_with_retries(method, url, *, max_retries=3, retry_delay=2, **kwargs):
    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            response = requests.request(method, url, **kwargs)

            if response.ok:
                return response

            last_error = RuntimeError(
                f"HTTP {response.status_code} on {url}\n{response.text}"
            )

            if attempt < max_retries:
                time.sleep(retry_delay * attempt)

        except requests.RequestException as exc:
            last_error = exc
            if attempt < max_retries:
                time.sleep(retry_delay * attempt)

    raise RuntimeError(f"Request failed after retries: {last_error}")


# ============================================================
# FILE EXTRACTION
# ============================================================

def extract_text_from_docx_path(docx_path: Path) -> str:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []

    with ZipFile(docx_path) as zf:
        with zf.open("word/document.xml") as f:
            tree = ET.parse(f)
            root = tree.getroot()

            for para in root.findall(".//w:p", ns):
                texts = []
                for node in para.findall(".//w:t", ns):
                    if node.text:
                        texts.append(node.text)
                para_text = "".join(texts).strip()
                if para_text:
                    paragraphs.append(para_text)

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError(f"No readable text found in DOCX: {docx_path.name}")
    return text


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = Path(tmp.name)

    try:
        return extract_text_from_docx_path(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    pages_text = []
    try:
        with pdfplumber.open(tmp_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"\n--- PAGE {i} ---\n{text.strip()}")
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    full_text = "\n".join(pages_text).strip()
    if not full_text:
        raise ValueError("No readable text found in uploaded PDF.")
    return full_text


def extract_text_from_uploaded_source(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    data = uploaded_file.getvalue()

    if suffix == ".pdf":
        return extract_text_from_pdf_bytes(data)

    if suffix == ".docx":
        return extract_text_from_docx_bytes(data)

    raise ValueError(f"Unsupported source file type: {uploaded_file.name}")


def extract_text_from_uploaded_prompt(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix != ".docx":
        raise ValueError(f"Prompt file must be DOCX: {uploaded_file.name}")
    return extract_text_from_docx_bytes(uploaded_file.getvalue())


def combine_source_files(source_files) -> str:
    all_blocks = []

    for uploaded in source_files:
        extracted = extract_text_from_uploaded_source(uploaded)
        all_blocks.append(f"\n\n===== FILE: {uploaded.name} =====\n{extracted}")

    raw_text = "\n".join(all_blocks).strip()
    if not raw_text:
        raise ValueError("No text extracted from uploaded source files.")
    return raw_text


# ============================================================
# DOCX CREATION
# ============================================================

def create_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        else:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# OPENAI HELPERS
# ============================================================

def openai_headers(config: AppConfig) -> dict:
    if not config.openai_api_key:
        raise ValueError("OpenAI API key is missing.")
    return {
        "Authorization": f"Bearer {config.openai_api_key}",
        "Content-Type": "application/json",
    }


def parse_openai_response(data: dict) -> str:
    if isinstance(data.get("output_text"), str) and data["output_text"].strip():
        return data["output_text"].strip()

    chunks = []

    for item in data.get("output", []):
        if not isinstance(item, dict):
            continue

        for part in item.get("content", []):
            if not isinstance(part, dict):
                continue

            part_type = part.get("type")
            if part_type in {"output_text", "text"}:
                txt = part.get("text", "")
                if txt:
                    chunks.append(txt)

    result = "\n".join(chunks).strip()
    if result:
        return result

    raise RuntimeError(f"Could not parse OpenAI response:\n{json.dumps(data, indent=2)}")


def call_openai(system_prompt: str, user_text: str, config: AppConfig) -> str:
    payload = {
        "model": config.openai_model,
        "instructions": system_prompt,
        "input": user_text,
    }

    response = request_with_retries(
        "POST",
        OPENAI_RESPONSES_URL,
        headers=openai_headers(config),
        json=payload,
        timeout=300,
        max_retries=3,
        retry_delay=2,
    )

    data = response.json()
    return parse_openai_response(data)


def generate_outline_with_gpt(raw_text: str, gpt_prompt_text: str, config: AppConfig) -> str:
    system_prompt = (
        "You are a financial document extraction assistant. "
        "Create one clean combined outline from the provided source text. "
        "Use only the source content. "
        "Do not add external facts, assumptions, predictions, recommendations, or calculations. "
        "If a required section is missing, write exactly: Not available in source data."
    )

    user_text = f"INSTRUCTIONS:\n{gpt_prompt_text}\n\nSOURCE TEXT:\n{raw_text}"
    return call_openai(system_prompt, user_text, config)


def generate_ppt_content_with_gpt(outline_text: str, ppt_prompt_text: str, config: AppConfig) -> str:
    system_prompt = (
        "You are a financial presentation content assistant. "
        "Convert the outline into slide-ready content. "
        "Use only the provided outline. "
        "Do not add external facts, assumptions, predictions, recommendations, or calculations. "
        "If a section is missing, write exactly: Not available in source data."
    )

    user_text = f"INSTRUCTIONS:\n{ppt_prompt_text}\n\nOUTLINE TEXT:\n{outline_text}"
    return call_openai(system_prompt, user_text, config)


# ============================================================
# GAMMA HELPERS
# ============================================================

def gamma_headers(config: AppConfig) -> dict:
    if not config.gamma_api_key or not config.gamma_api_key.startswith("sk-gamma-"):
        raise ValueError("Gamma API key is missing or invalid.")
    return {
        "X-API-KEY": config.gamma_api_key,
        "Content-Type": "application/json",
        "Accept": "application/json",
    }


def gamma_post_json(url: str, payload: dict, config: AppConfig) -> dict:
    response = request_with_retries(
        "POST",
        url,
        headers=gamma_headers(config),
        json=payload,
        timeout=300,
        max_retries=3,
        retry_delay=2,
    )
    return response.json()


def gamma_get_json(url: str, config: AppConfig) -> dict:
    response = request_with_retries(
        "GET",
        url,
        headers=gamma_headers(config),
        timeout=300,
        max_retries=3,
        retry_delay=2,
    )
    return response.json()


def create_gamma_from_template(ppt_content: str, config: AppConfig) -> str:
    payload = {
        "gammaId": config.gamma_id,
        "prompt": ppt_content,
        "themeId": config.gamma_theme_id,
        "exportAs": "pptx",
        "imageOptions": {
            "model": "imagen-4-pro",
            "style": "photorealistic",
        },
    }

    if config.gamma_folder_ids:
        payload["folderIds"] = config.gamma_folder_ids

    resp = gamma_post_json(GAMMA_FROM_TEMPLATE_URL, payload, config)
    generation_id = resp.get("generationId") or resp.get("id")

    if not generation_id:
        raise RuntimeError(f"generationId missing in Gamma response:\n{json.dumps(resp, indent=2)}")

    return generation_id


def wait_for_gamma_completion(
    generation_id: str,
    config: AppConfig,
    poll_interval: int = 5,
    timeout_seconds: int = 900,
) -> dict:
    deadline = time.time() + timeout_seconds
    status_url = GAMMA_STATUS_URL.format(generation_id=generation_id)

    while time.time() < deadline:
        resp = gamma_get_json(status_url, config)
        status = str(resp.get("status", "")).lower()

        if status == "completed":
            return resp

        if status in {"failed", "error", "cancelled"}:
            raise RuntimeError(f"Gamma generation failed:\n{json.dumps(resp, indent=2)}")

        time.sleep(poll_interval)

    raise TimeoutError("Timed out waiting for Gamma generation to complete.")


def find_gamma_link(status_resp: dict):
    candidates = []

    def walk(obj):
        if isinstance(obj, dict):
            for _, value in obj.items():
                if isinstance(value, str) and value.startswith("http") and "gamma" in value.lower():
                    candidates.append(value)
                else:
                    walk(value)
        elif isinstance(obj, list):
            for item in obj:
                walk(item)

    walk(status_resp)
    return candidates[0] if candidates else None